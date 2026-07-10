#!/usr/bin/env python3
"""
Gera o templates/default/reference.docx usado pelo Markforge como
"reference-doc" do Pandoc. Define:
  - Estilos de Titulo/Subtitulo/Autor/Data (capa)
  - Estilos de Heading 1-4, corpo, citacao e codigo
  - Cabecalho e rodape (com numero de pagina) em todas as paginas,
    exceto a capa (different first page header/footer)

Roda localmente com python-docx (sem rede). O resultado binario e'
commitado no repo em templates/default/reference.docx.
"""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # azul-marinho "Markforge"
ACCENT_HEX = '1F3A5F'
MUTED = RGBColor(0x6B, 0x72, 0x80)    # cinza texto secundario
FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"


def set_font(style, name=FONT_BODY, size=11, bold=False, italic=False, color=None):
    f = style.font
    f.name = name
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    if color is not None:
        f.color.rgb = color
    # garante que fontes do East Asian / complex script tambem sigam o nome
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)


def add_field(paragraph, field_code):
    """Insere um campo de Word (ex: PAGE, NUMPAGES) num paragrafo."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


def build():
    doc = Document()

    styles = doc.styles

    # ---- Normal / corpo ----
    normal = styles['Normal']
    set_font(normal, FONT_BODY, 11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    # ---- Titulo da capa ----
    if 'Title' in [s.name for s in styles]:
        title = styles['Title']
    else:
        title = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    set_font(title, FONT_HEAD, 30, bold=True, color=ACCENT)
    title.paragraph_format.space_before = Pt(180)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if 'Subtitle' in [s.name for s in styles]:
        subtitle = styles['Subtitle']
    else:
        subtitle = styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    set_font(subtitle, FONT_BODY, 16, italic=True, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 'Author' e 'Date' sao estilos que o Pandoc usa para blocos de metadados
    for name, size in (('Author', 13), ('Date', 12)):
        if name in [s.name for s in styles]:
            st = styles[name]
        else:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_font(st, FONT_BODY, size, color=MUTED)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.space_after = Pt(4)

    # ---- Headings ----
    heading_sizes = {1: 22, 2: 17, 3: 14, 4: 12}
    for level, size in heading_sizes.items():
        st = styles[f'Heading {level}']
        set_font(st, FONT_HEAD, size, bold=True, color=ACCENT)
        st.paragraph_format.space_before = Pt(20 if level == 1 else 14)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.keep_with_next = True

    # ---- Citacao ----
    if 'Quote' in [s.name for s in styles]:
        quote = styles['Quote']
    else:
        quote = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    set_font(quote, FONT_BODY, 11, italic=True, color=MUTED)
    quote.paragraph_format.left_indent = Cm(1)
    quote.paragraph_format.space_after = Pt(8)

    # ---- Codigo ----
    code_style_name = 'Source Code'
    if code_style_name not in [s.name for s in styles]:
        code = styles.add_style(code_style_name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles[code_style_name]
    set_font(code, 'Consolas', 10, color=RGBColor(0x2B, 0x2B, 0x2B))
    code.paragraph_format.left_indent = Cm(0.5)
    code.paragraph_format.space_after = Pt(8)

    if 'Verbatim Char' in [s.name for s in styles]:
        vch = styles['Verbatim Char']
    else:
        vch = styles.add_style('Verbatim Char', WD_STYLE_TYPE.CHARACTER)
    set_font(vch, 'Consolas', 10, color=RGBColor(0x2B, 0x2B, 0x2B))

    # ---- 'Compact' (usado pelo Pandoc em celulas de tabela e listas apertadas) ----
    if 'Compact' in [s.name for s in styles]:
        compact = styles['Compact']
    else:
        compact = styles.add_style('Compact', WD_STYLE_TYPE.PARAGRAPH)
    set_font(compact, FONT_BODY, 10)
    compact.paragraph_format.space_after = Pt(2)
    compact.paragraph_format.line_spacing = 1.0

    # ---- Tabelas: estilo 'Table' (nome que o Pandoc procura no reference-doc) ----
    if 'Table' in [s.name for s in styles]:
        table_style = styles['Table']
    else:
        table_style = styles.add_style('Table', WD_STYLE_TYPE.TABLE)
    set_font(table_style, FONT_BODY, 10)
    tbl_el = table_style.element
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.append(tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'D0D4DA')
        borders.append(el)
    tblPr.append(borders)
    cellMar = OxmlElement('w:tblCellMar')
    for edge, val in (('top', '40'), ('bottom', '40'), ('left', '80'), ('right', '80')):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), val)
        el.set(qn('w:type'), 'dxa')
        cellMar.append(el)
    tblPr.append(cellMar)
    # linha de cabecalho da tabela com fundo suave e negrito
    table_style.element.get_or_add_rPr()  # garante rPr geral
    # w:tblStylePr para a primeira linha (firstRow) com shading + negrito
    tblStylePr = OxmlElement('w:tblStylePr')
    tblStylePr.set(qn('w:type'), 'firstRow')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF')
    rPr.append(color)
    tblStylePr.append(rPr)
    tcPr = OxmlElement('w:tcPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), ACCENT_HEX)
    tcPr.append(shd)
    tblStylePr.append(tcPr)
    tbl_el.append(tblStylePr)

    # ---- Margens ----
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.different_first_page_header_footer = True

    # ---- Cabecalho (paginas 2+) ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ''
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run('MARKFORGE')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.bold = True
    run.font.name = FONT_BODY
    # linha fina abaixo do cabecalho
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'D0D4DA')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ---- Rodape (paginas 2+): "Pagina X de Y" ----
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run('Página ')
    r1.font.size = Pt(8)
    r1.font.color.rgb = MUTED
    r1.font.name = FONT_BODY
    add_field(fp, 'PAGE')
    r2 = fp.add_run(' de ')
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED
    r2.font.name = FONT_BODY
    add_field(fp, 'NUMPAGES')
    for r in fp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
        r.font.name = FONT_BODY

    # ---- Capa: cabecalho/rodape em branco (different_first_page) ----
    fh = section.first_page_header
    fh.is_linked_to_previous = False
    fh.paragraphs[0].text = ''
    ff = section.first_page_footer
    ff.is_linked_to_previous = False
    ff.paragraphs[0].text = ''

    # paragrafo de amostra (sera substituido pelo conteudo real do pandoc)
    doc.add_paragraph('Corpo do documento.', style='Normal')

    doc.save('templates/default/reference.docx')
    print('OK: templates/default/reference.docx gerado')


if __name__ == '__main__':
    build()
